"""Generate SUNDAE Project Documentation as Word (.docx) file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Style setup --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ============================================================
# COVER PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SUNDAE")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Enterprise AI Chatbot Platform")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Project Documentation")
run.font.size = Pt(14)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Tech Stack: ").bold = True
meta.add_run("FastAPI + React + Supabase + Ollama")

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta2.add_run("Version: ").bold = True
meta2.add_run("March 2026 (Updated)")

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Project Overview",
    "2. Architecture",
    "3. Backend (FastAPI + Python)",
    "   3.1 Project Structure",
    "   3.2 Database Schema",
    "   3.3 AI Services",
    "   3.4 Auth Middleware",
    "   3.5 API Endpoints",
    "4. Frontend (React + Vite + Tailwind v4)",
    "   4.1 Project Structure",
    "   4.2 Token Protection Strategy",
    "5. Supabase Auth Integration",
    "   5.1 Auth Flow",
    "   5.2 Registration Flow",
    "6. RLS Security (Row Level Security)",
    "7. SQL Migrations (001-009)",
    "8. Bug Fixes Summary",
    "9. LLM Model Configuration",
    "10. Deployment & Run Commands",
    "11. Next Steps & Roadmap",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. PROJECT OVERVIEW
# ============================================================
doc.add_heading("1. Project Overview", level=1)
doc.add_paragraph(
    "SUNDAE is an Enterprise AI Chatbot Platform designed for Thai-language customer support. "
    "It uses a Retrieval-Augmented Generation (RAG) architecture to provide accurate, "
    "context-aware responses based on uploaded knowledge base documents."
)
doc.add_paragraph(
    "The system supports multi-tenant organizations, role-based access control (Admin, Support, User), "
    "omnichannel messaging (Web + LINE), human handoff capability, and streaming AI responses."
)

doc.add_heading("Key Features", level=2)
features = [
    "RAG-based AI chatbot with Thai language support",
    "Document upload & intelligent chunking (Parent-Child strategy)",
    "Vector search with BGE-M3 embeddings (1024 dimensions)",
    "Reranking with BGE-Reranker-v2-M3 for improved accuracy",
    "Streaming SSE responses with Ollama/Qwen2.5",
    "Multi-tenant organization support",
    "Role-based access: Admin, Support, User",
    "User approval workflow",
    "Human handoff (AI to human agent)",
    "Omnichannel: Web Chat + LINE (planned)",
    "JWT authentication via Supabase Auth",
    "Row Level Security (RLS) on all tables",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("Tech Stack", level=2)
add_table(
    ["Layer", "Technology", "Details"],
    [
        ["Backend", "FastAPI (Python)", "REST API + SSE streaming"],
        ["Frontend", "React + Vite + Tailwind v4", "TypeScript SPA"],
        ["Database", "Supabase PostgreSQL + pgvector", "Vector search + RLS"],
        ["Auth", "Supabase Auth", "JWT-based authentication"],
        ["Embedding", "BAAI/bge-m3", "1024-dim multilingual embeddings"],
        ["Reranker", "BAAI/bge-reranker-v2-m3", "Cross-encoder reranking"],
        ["LLM", "Ollama / Qwen2.5:3b", "Local LLM inference"],
        ["Containerization", "Docker + Docker Compose", "Multi-service deployment"],
    ],
)

doc.add_page_break()

# ============================================================
# 2. ARCHITECTURE
# ============================================================
doc.add_heading("2. Architecture", level=1)
doc.add_paragraph(
    "The system follows a three-tier architecture with clear separation between "
    "Frontend, Backend, and Database layers."
)

doc.add_heading("Architecture Diagram (Text)", level=2)
arch_text = """
+---------------------------+       +-------------------+
|   Frontend (React+Vite)   |       |   Supabase Auth   |
|  - LoginPage              | <---> |  - Auth Service   |
|  - Dashboard              |  JWT  |  - user_profiles  |
|  - Knowledge Base         |       |  - RLS Policies   |
|  - Bots                   |       +-------------------+
|  - Inbox                  |
|  - Web Chat               |
|  - Approvals              |
+------------+--------------+
             | Axios + Bearer Token
             v
+---------------------------+
|   Backend (FastAPI)       |
|  - Document Router        |
|  - Chat Router            |
|  - Inbox Router           |
|  - Bot Router             |
|  - Chunking Service       |
|  - AI Models (BGE-M3)     |
|  - Vector Search          |
|  - LLM Generator (Ollama) |
+------------+--------------+
             |
             v
+---------------------------+
|  Supabase PostgreSQL      |
|  - organizations          |
|  - bots                   |
|  - documents              |
|  - parent_chunks          |
|  - child_chunks (pgvector)|
|  - chat_sessions          |
|  - chat_messages          |
+---------------------------+
"""
add_code(arch_text.strip())

doc.add_heading("Data Flow: Chat Query", level=2)
chat_flow = [
    "1. User sends message via Web Chat or LINE",
    "2. Frontend sends POST request with JWT token to Backend",
    "3. Backend authenticates user via Supabase JWT",
    "4. Message is embedded using BGE-M3 model (1024 dims)",
    "5. Vector search finds top-K matching child chunks via pgvector",
    "6. Results are reranked using BGE-Reranker-v2-M3",
    "7. Top chunks' parent content is retrieved as context",
    "8. Context + message sent to Ollama/Qwen2.5 LLM",
    "9. LLM generates response (streamed via SSE)",
    "10. Response delivered to user in real-time",
]
for step in chat_flow:
    doc.add_paragraph(step, style="List Bullet")

doc.add_heading("Data Flow: Document Upload", level=2)
upload_flow = [
    "1. Admin uploads document (PDF/DOCX) via Knowledge Base page",
    "2. Backend extracts text content",
    "3. Thai Text Splitter chunks content into parent chunks (~800 chars)",
    "4. Each parent chunk is further split into child chunks (~200 chars)",
    "5. Child chunks are embedded using BGE-M3 (1024-dim vectors)",
    "6. All chunks + embeddings stored in Supabase PostgreSQL",
]
for step in upload_flow:
    doc.add_paragraph(step, style="List Bullet")

doc.add_page_break()

# ============================================================
# 3. BACKEND
# ============================================================
doc.add_heading("3. Backend (FastAPI + Python)", level=1)

doc.add_heading("3.1 Project Structure", level=2)
structure = """backend/
+-- app/
|   +-- main.py              # FastAPI app + CORS
|   +-- core/
|   |   +-- config.py         # Settings from .env
|   |   +-- auth.py           # JWT middleware
|   |   +-- database.py       # Supabase async client
|   +-- routers/
|   |   +-- document.py       # Upload, list, delete
|   |   +-- chat.py           # Omnichannel chat + SSE
|   |   +-- inbox.py          # Session mgmt + handoff
|   |   +-- bot.py            # Bot CRUD
|   |   +-- health.py         # Health check
|   +-- services/
|   |   +-- chunking.py       # Thai text splitter
|   |   +-- ai_models.py      # Embedding + Reranker
|   |   +-- vector_search.py  # Supabase RPC search
|   |   +-- llm_generator.py  # Ollama/Qwen2.5
|   +-- models/               # Pydantic schemas
+-- sql/                      # DB migrations (001-009)
+-- tests/                    # Unit tests
+-- requirements.txt
+-- Dockerfile"""
add_code(structure)

doc.add_heading("3.2 Database Schema", level=2)
add_table(
    ["Table", "Primary Key", "Description"],
    [
        ["organizations", "UUID", "Multi-tenant root organization"],
        ["user_profiles", "UUID (FK -> auth.users)", "role, is_approved, email, full_name"],
        ["bots", "UUID", "Bot config: prompt, line_access_token, is_web_enabled"],
        ["documents", "UUID", "Uploaded files: file_path, status, FK -> bots"],
        ["document_parent_chunks", "UUID", "Large content chunks + metadata"],
        ["document_child_chunks", "UUID", "Small chunks + embedding vector(1024)"],
        ["chat_sessions", "UUID", "Chat sessions: platform_source, status"],
        ["chat_messages", "UUID", "Messages: role (user/assistant/system/admin), content"],
    ],
)
doc.add_paragraph(
    "RPC Function: match_child_chunks - performs cosine similarity search on pgvector "
    "with optional bot_id filtering (added in migration 006)."
)

doc.add_heading("3.3 AI Services", level=2)
add_table(
    ["Service", "Model", "Purpose"],
    [
        ["Embedding", "BAAI/bge-m3 (1024 dims)", "Convert text to vector embeddings"],
        ["Reranker", "BAAI/bge-reranker-v2-m3", "Re-rank search results for relevance"],
        ["LLM", "Ollama/qwen2.5:3b", "Generate answers from context (configurable)"],
        ["Chunking", "Custom Thai Splitter", "Split text into parent/child chunks"],
    ],
)

doc.add_heading("3.4 Auth Middleware (core/auth.py)", level=2)
doc.add_paragraph("Three levels of authentication/authorization:")
add_table(
    ["Function", "Purpose"],
    [
        ["get_current_user", "Verify Bearer token, fetch user_profiles from DB (service role)"],
        ["require_approved", "Check is_approved = true, return 403 if not"],
        ["require_role(...)", "Check role + is_approved, return 403 if unauthorized"],
    ],
)
doc.add_paragraph(
    "Backend uses Supabase Service Role Key, which bypasses RLS entirely. "
    "This ensures the backend always reads the real is_approved value from DB."
)

doc.add_heading("3.5 API Endpoints", level=2)

doc.add_heading("Document Router (/documents)", level=3)
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/documents/upload", "Admin/Support", "Upload document (PDF/DOCX)"],
        ["GET", "/documents/", "Approved user", "List documents for organization"],
        ["DELETE", "/documents/{id}", "Admin/Support", "Delete document and its chunks"],
    ],
)

doc.add_heading("Chat Router (/chat)", level=3)
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["POST", "/chat/ask", "Approved user", "Send message, get AI response"],
        ["POST", "/chat/ask-stream", "Approved user", "Send message, get streaming SSE response"],
    ],
)

doc.add_heading("Inbox Router (/inbox)", level=3)
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["GET", "/inbox/sessions", "Admin/Support", "List chat sessions"],
        ["GET", "/inbox/sessions/{id}/messages", "Admin/Support", "Get messages for session"],
        ["POST", "/inbox/sessions/{id}/reply", "Admin/Support", "Send admin reply (human handoff)"],
        ["PATCH", "/inbox/sessions/{id}/status", "Admin/Support", "Update session status"],
    ],
)

doc.add_heading("Bot Router (/bots)", level=3)
add_table(
    ["Method", "Endpoint", "Auth", "Description"],
    [
        ["GET", "/bots/", "Approved user", "List bots for organization"],
        ["POST", "/bots/", "Admin/Support", "Create new bot"],
        ["PUT", "/bots/{id}", "Admin/Support", "Update bot settings"],
        ["DELETE", "/bots/{id}", "Admin", "Delete bot"],
    ],
)

doc.add_page_break()

# ============================================================
# 4. FRONTEND
# ============================================================
doc.add_heading("4. Frontend (React + Vite + Tailwind v4)", level=1)

doc.add_heading("4.1 Project Structure", level=2)
fe_structure = """frontend/src/
+-- api/
|   +-- supabaseClient.ts    # Singleton Supabase client
|   +-- axios.ts             # JWT interceptor (3 layers)
|   +-- endpoints.ts         # API calls
+-- store/
|   +-- authStore.ts         # Zustand (signIn/signOut/fetchProfile)
|   +-- toastStore.ts        # Toast notification state
+-- types/
|   +-- index.ts             # TypeScript interfaces
+-- components/
|   +-- ProtectedRoute.tsx   # Role guard
|   +-- ToastContainer.tsx   # Global toast UI
+-- layouts/
|   +-- DashboardLayout.tsx  # Sidebar + approval lockout
|   +-- AuthLayout.tsx       # Login background
+-- pages/
|   +-- LoginPage.tsx        # Login + Registration tabs
|   +-- DashboardPage.tsx    # 3 role-based states
|   +-- WebChatPage.tsx      # Chat + streaming + cancel
|   +-- ApprovalsPage.tsx    # Admin/Support approval list
|   +-- KnowledgeBasePage.tsx
|   +-- BotsPage.tsx
|   +-- InboxPage.tsx        # Human handoff inbox
|   +-- IntegrationPage.tsx  # LINE integration (mock)
+-- App.tsx                  # AuthProvider + Routing
+-- index.css                # NT CI Design System
+-- main.tsx                 # Entry point"""
add_code(fe_structure)

doc.add_heading("4.2 Token Protection Strategy (axios.ts)", level=2)
doc.add_paragraph("Three layers protect against expired JWT tokens:")

add_table(
    ["Layer", "Location", "Mechanism"],
    [
        ["Layer 1", "Request Interceptor", "getValidToken(): check expiry, refresh if < 5 min remaining"],
        ["Layer 2", "Response Interceptor", "On 401: refreshTokenOnce() -> retry request once"],
        ["Layer 3", "supabaseClient.ts", "Periodic refresh every 4 min + refresh on tab re-focus after 5 min"],
    ],
)
doc.add_paragraph(
    "Mutex: refreshPromise prevents concurrent refresh calls that would invalidate the refresh token."
)

doc.add_page_break()

# ============================================================
# 5. SUPABASE AUTH
# ============================================================
doc.add_heading("5. Supabase Auth Integration", level=1)

doc.add_heading("5.1 Login Flow", level=2)
login_flow = """1. User enters email/password
2. Frontend calls supabase.auth.signInWithPassword()
3. Supabase Auth returns Session + JWT
4. Frontend queries user_profiles (role, is_approved, email, full_name)
5. Store state in Zustand
6. If is_approved = true -> Show Dashboard
   If is_approved = false -> Show Lockout Screen"""
add_code(login_flow)

doc.add_heading("5.2 Registration Flow", level=2)
reg_flow = """1. User fills registration form (email, password, full_name)
2. Frontend calls supabase.auth.signUp()
3. Supabase creates auth user
4. DB trigger auto-creates user_profiles entry:
   - role = 'user'
   - is_approved = false
   - organization = default org
5. User must login again + wait for Admin/Support approval"""
add_code(reg_flow)

# ============================================================
# 6. RLS SECURITY
# ============================================================
doc.add_heading("6. RLS Security (Row Level Security)", level=1)
doc.add_paragraph(
    "All tables use Row Level Security (RLS) policies to enforce data access rules at the database level."
)

doc.add_heading("Helper Function", level=2)
add_code("""get_my_role() RETURNS TEXT -- SECURITY DEFINER (bypasses RLS)
  -> SELECT role FROM user_profiles WHERE id = auth.uid()
  -> Prevents infinite recursion in RLS policies""")

doc.add_heading("user_profiles Policies", level=2)
add_table(
    ["Operation", "Policy Rule"],
    [
        ["SELECT", "Own profile (id = auth.uid()) OR Support/Admin can read all"],
        ["UPDATE", "Only Support/Admin (USING + WITH CHECK)"],
        ["INSERT", "Only own profile (id = auth.uid())"],
    ],
)

doc.add_heading("organizations Policies", level=2)
add_table(
    ["Operation", "Policy Rule"],
    [
        ["SELECT", "User can read own org (via user_profiles.organization_id)"],
        ["ALL", "Service role has full access"],
    ],
)

doc.add_page_break()

# ============================================================
# 7. SQL MIGRATIONS
# ============================================================
doc.add_heading("7. SQL Migrations (001-009)", level=1)
add_table(
    ["#", "Migration", "Status", "Purpose"],
    [
        ["001", "Main Schema", "Completed", "Create all tables + match_child_chunks RPC"],
        ["002", "Add Columns", "Completed", "LINE integration + session status fields"],
        ["003", "user_profiles + RLS", "Completed", "Role system + is_approved + get_my_role()"],
        ["004", "Auth Trigger", "Completed", "Auto-create profile on user signup"],
        ["005", "Support Account", "Completed", "Create support user account"],
        ["006", "Bot Filter", "Completed", "Vector search filters by bot_id"],
        ["007", "Admin Role", "Completed", "Allow 'admin' role in chat_messages"],
        ["008", "Fix Org RLS", "Completed", "Fix 406 error on organizations table"],
        ["009", "Fix UPDATE RLS", "Completed", "Add WITH CHECK to UPDATE policy"],
    ],
)

# ============================================================
# 8. BUG FIXES
# ============================================================
doc.add_heading("8. Bug Fixes Summary", level=1)
add_table(
    ["#", "Bug", "Root Cause", "Fix", "File"],
    [
        ["B1", "Loading screen stuck", "No .catch() + no timeout", "Added .catch() + 5s timeout", "App.tsx"],
        ["B2", "Missing email/full_name", "Missing DB columns", "Added columns + sync types", "002 migration"],
        ["B3", "RLS blocks profile read", "No SELECT for self", "Added id = auth.uid()", "003 migration"],
        ["B4", "RLS infinite recursion", "Subquery reads user_profiles", "Created get_my_role() SECURITY DEFINER", "003 migration"],
        ["B5", "Stream not firing", "Dynamic import hangs in Vite", "Changed to static import", "endpoints.ts"],
        ["B6", "Organizations 406", "RLS uses missing JWT claim", "New policy with subquery", "008 migration"],
        ["B7", "Login 'Failed to fetch'", "Missing catch block", "Added catch + Thai error msg", "authStore.ts"],
        ["B8", "Approve not working", "Missing WITH CHECK + no verify", "Added WITH CHECK + select()", "009 migration"],
        ["B9", "Stream fails after time", "Expired token not refreshed", "Use getValidToken()", "endpoints.ts"],
        ["B10", "Ollama unloads model", "keep_alive too short (30m)", "Changed to keep_alive: 4h", "llm_generator.py"],
        ["B11", "Inbox layout wrong", "User perspective instead of admin", "Switched to admin perspective", "InboxPage.tsx"],
    ],
)

doc.add_page_break()

# ============================================================
# 9. LLM MODEL
# ============================================================
doc.add_heading("9. LLM Model Configuration", level=1)

doc.add_heading("Current Model", level=2)
add_table(
    ["Model", "RAM Required", "Status"],
    [
        ["qwen3:14b", "~16 GB", "Not available (insufficient RAM)"],
        ["qwen2.5:7b", "~8 GB", "Available if RAM permits"],
        ["qwen2.5:3b (current)", "~4 GB", "Currently in use"],
    ],
)

doc.add_paragraph("Model is configured via LLM_MODEL environment variable in backend/.env")

doc.add_heading("Quality Comparison", level=2)
add_table(
    ["Aspect", "qwen3:14b", "qwen2.5:3b"],
    [
        ["Thai language", "Excellent", "Good (may be shorter)"],
        ["Accuracy", "High", "Medium-High"],
        ["Speed", "Slower", "Faster"],
        ["RAM", "16 GB", "4 GB"],
    ],
)

doc.add_paragraph(
    "To upgrade: change only LLM_MODEL in backend/.env. No code changes needed."
)

# ============================================================
# 10. DEPLOYMENT
# ============================================================
doc.add_heading("10. Deployment & Run Commands", level=1)

doc.add_heading("Development", level=2)
add_code("""# Backend
cd backend
uvicorn app.main:app --host 0.0.0.0 --port 8001 --reload

# Frontend
cd frontend
npm run dev""")

doc.add_heading("Docker", level=2)
add_code("""# Build and run all services
docker-compose up --build

# Services:
#   - backend:  port 8001
#   - frontend: port 80 (nginx)
#   - ollama:   port 11434""")

doc.add_heading("Environment Variables", level=2)

doc.add_heading("Backend (.env)", level=3)
add_table(
    ["Variable", "Description", "Example"],
    [
        ["SUPABASE_URL", "Supabase project URL", "https://xxx.supabase.co"],
        ["SUPABASE_SERVICE_KEY", "Service role key (bypasses RLS)", "eyJ..."],
        ["SUPABASE_ANON_KEY", "Anonymous key", "eyJ..."],
        ["LLM_MODEL", "Ollama model name", "qwen2.5:3b"],
        ["OLLAMA_BASE_URL", "Ollama API URL", "http://localhost:11434"],
        ["LINE_CHANNEL_SECRET", "LINE webhook signature verify", "(future)"],
    ],
)

doc.add_heading("Frontend (.env)", level=3)
add_table(
    ["Variable", "Description"],
    [
        ["VITE_SUPABASE_URL", "Supabase project URL"],
        ["VITE_SUPABASE_ANON_KEY", "Supabase anonymous key"],
        ["VITE_API_BASE_URL", "Backend API URL (e.g., http://localhost:8001)"],
    ],
)

doc.add_heading("Test Accounts", level=2)
add_table(
    ["Email", "Password", "Role", "Approved"],
    [
        ["admin@sundae.local", "Sundae@2025", "admin", "Yes"],
        ["support@sundae.local", "Sundae@2025", "support", "Yes"],
    ],
)

# ============================================================
# 11. NEXT STEPS
# ============================================================
doc.add_heading("11. Next Steps & Roadmap", level=1)
add_table(
    ["#", "Task", "Details"],
    [
        ["1", "Test keep_alive 4h", "Verify streaming works after 30+ min idle"],
        ["2", "LINE Webhook", "Create POST /webhook/line endpoint + LINE reply service"],
        ["3", "Docker deployment", "Test full build + run on Docker for production"],
        ["4", "User profile edit", "Allow users to update their own full_name"],
        ["5", "Dark mode", "Add theme switcher"],
    ],
)

doc.add_heading("LINE Webhook Implementation Plan", level=2)
line_steps = [
    "1. Create backend/app/routers/line_webhook.py - POST /webhook/line endpoint",
    "2. Verify signature with HMAC-SHA256 + LINE_CHANNEL_SECRET",
    "3. Parse LINE events: extract user_id, message text, reply_token",
    "4. Match bot by line_access_token in DB",
    "5. Run existing RAG pipeline (same as chat.py)",
    "6. Reply via LINE Messaging API with reply_token",
    "7. Create backend/app/services/line_service.py for LINE API calls",
    "8. Connect IntegrationPage.tsx to real Bot API (replace mock useState)",
]
for step in line_steps:
    doc.add_paragraph(step, style="List Bullet")

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), "SUNDAE_Project_Documentation.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
