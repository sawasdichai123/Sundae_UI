"""
Generate SUNDAE 5-Chapter Academic Project Report (.docx)
Format: Bangkok University Senior Project Report Style
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# STYLE SETUP — TH Sarabun New style (fallback to Cordia New / Calibri)
# ============================================================
FONT_TH = "TH Sarabun New"
FONT_EN = "TH Sarabun New"
FONT_SIZE = Pt(16)  # 16pt is standard for TH Sarabun in Thai academic
FONT_SIZE_HEADING = Pt(18)
FONT_SIZE_TITLE = Pt(20)
LINE_SPACING = 1.5

style_normal = doc.styles["Normal"]
style_normal.font.name = FONT_TH
style_normal.font.size = FONT_SIZE
style_normal.paragraph_format.line_spacing = LINE_SPACING

# Set font for East Asian / Thai
rFonts = style_normal.element.rPr.rFonts if style_normal.element.rPr is not None else None

def set_paragraph(text, bold=False, size=None, align=None, space_after=None, space_before=None, indent_left=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if indent_left is not None:
        p.paragraph_format.left_indent = Cm(indent_left)
    run = p.add_run(text)
    run.font.name = FONT_TH
    run.font.size = size or FONT_SIZE
    run.bold = bold
    return p

def set_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = FONT_TH
        run.font.size = FONT_SIZE_HEADING if level <= 2 else FONT_SIZE
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = FONT_TH
                r.font.size = Pt(14)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT_TH
                    r.font.size = Pt(14)
    doc.add_paragraph()
    return table

def add_empty_lines(n=1):
    for _ in range(n):
        doc.add_paragraph()

# ============================================================
# ปกหน้า (Cover Page)
# ============================================================
add_empty_lines(3)

set_paragraph(
    "แพลตฟอร์มแชทบอท AI สำหรับองค์กร SUNDAE",
    bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER
)
set_paragraph(
    "SUNDAE — Enterprise AI Chatbot Platform\nwith Retrieval-Augmented Generation",
    bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER
)

add_empty_lines(2)

set_paragraph("โดย", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)

add_empty_lines(1)

set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 1]", bold=False, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("รหัส [XXXXXXXXXX]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(0)
set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 2]", bold=False, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("รหัส [XXXXXXXXXX]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(0)
set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 3]", bold=False, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("รหัส [XXXXXXXXXX]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)

add_empty_lines(1)

set_paragraph("อาจารย์ที่ปรึกษา", bold=True, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("[ตำแหน่งทางวิชาการ + ชื่อ-นามสกุล อาจารย์ที่ปรึกษา]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)

add_empty_lines(1)

set_paragraph("หลักสูตรวิทยาศาสตรบัณฑิต", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("ภาคการศึกษาที่ [X] ปีการศึกษา [25XX]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("ภาควิชา[ชื่อภาควิชา] คณะ[ชื่อคณะ]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("[ชื่อมหาวิทยาลัย]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# ใบรับรอง (Certification Page)
# ============================================================
set_paragraph("ใบรับรองวิชาโครงงาน", bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
set_paragraph("เรื่อง", bold=True, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(0)
set_paragraph(
    "แพลตฟอร์มแชทบอท AI สำหรับองค์กร SUNDAE\n"
    "SUNDAE — Enterprise AI Chatbot Platform\n"
    "with Retrieval-Augmented Generation",
    bold=True, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER
)
add_empty_lines(0)
set_paragraph("โดย", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(0)
set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 1]   รหัส [XXXXXXXXXX]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 2]   รหัส [XXXXXXXXXX]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 3]   รหัส [XXXXXXXXXX]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)

add_empty_lines(1)

set_paragraph(
    "รายงานฉบับนี้ได้รับการตรวจสอบและอนุมัติให้เป็นส่วนหนึ่งของ\n"
    "วิชาโครงงาน หลักสูตรวิทยาศาสตรบัณฑิต\n"
    "ภาคการศึกษาที่ [X] ปีการศึกษา [25XX]",
    size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER
)

add_empty_lines(1)

set_paragraph("...................................................... อาจารย์ที่ปรึกษาโครงงาน", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("([ชื่อ-นามสกุล อาจารย์ที่ปรึกษา])", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(0)
set_paragraph("...................................................... กรรมการสอบปากเปล่า", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("([ชื่อ-นามสกุล กรรมการ 1])", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
add_empty_lines(0)
set_paragraph("...................................................... กรรมการสอบปากเปล่า", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)
set_paragraph("([ชื่อ-นามสกุล กรรมการ 2])", size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# บทคัดย่อ (Thai Abstract)
# ============================================================
set_paragraph("บทคัดย่อ", bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# Metadata table
p = doc.add_paragraph()
run = p.add_run("ชื่อหัวข้อ\t\t")
run.bold = True
run.font.name = FONT_TH
run.font.size = FONT_SIZE
run2 = p.add_run("แพลตฟอร์มแชทบอท AI สำหรับองค์กร SUNDAE")
run2.font.name = FONT_TH
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("ผู้ร่วมโครงการ\t\t")
run.bold = True
run.font.name = FONT_TH
run.font.size = FONT_SIZE
run2 = p.add_run("[ชื่อ-นามสกุล 1]  [รหัส]\n\t\t\t[ชื่อ-นามสกุล 2]  [รหัส]\n\t\t\t[ชื่อ-นามสกุล 3]  [รหัส]")
run2.font.name = FONT_TH
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("อาจารย์ที่ปรึกษา\t")
run.bold = True
run.font.name = FONT_TH
run.font.size = FONT_SIZE
run2 = p.add_run("[ตำแหน่งทางวิชาการ + ชื่อ-นามสกุล]")
run2.font.name = FONT_TH
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("ระดับการศึกษา\t\t")
run.bold = True
run.font.name = FONT_TH
run.font.size = FONT_SIZE
run2 = p.add_run("วิทยาศาสตรบัณฑิต")
run2.font.name = FONT_TH
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("ภาควิชา\t\t\t")
run.bold = True
run.font.name = FONT_TH
run.font.size = FONT_SIZE
run2 = p.add_run("[ชื่อภาควิชา]")
run2.font.name = FONT_TH
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("ปีการศึกษา\t\t")
run.bold = True
run.font.name = FONT_TH
run.font.size = FONT_SIZE
run2 = p.add_run("[25XX]")
run2.font.name = FONT_TH
run2.font.size = FONT_SIZE

add_empty_lines(1)

abstract_th = (
    "\tแพลตฟอร์มแชทบอท AI สำหรับองค์กร SUNDAE (SUNDAE — Enterprise AI Chatbot Platform "
    "with Retrieval-Augmented Generation) จัดทำขึ้นเพื่อพัฒนาระบบแชทบอทอัจฉริยะที่สามารถ"
    "ตอบคำถามได้อย่างแม่นยำโดยอ้างอิงจากฐานความรู้ขององค์กร โดยใช้เทคนิค Retrieval-Augmented "
    "Generation (RAG) ที่ผสมผสานการค้นหาข้อมูลเชิงความหมาย (Semantic Search) ด้วย Vector "
    "Embedding และการสร้างคำตอบด้วย Large Language Model (LLM) เข้าด้วยกัน "
    "ระบบรองรับภาษาไทยเป็นหลัก สามารถอัปโหลดเอกสาร (PDF/DOCX) เพื่อสร้างฐานความรู้ "
    "มีระบบจัดการผู้ใช้แบบ Multi-tenant พร้อม Role-Based Access Control (Admin, Support, User) "
    "รองรับการสนทนาแบบ Omnichannel ทั้ง Web Chat และ LINE "
    "และมีระบบ Human Handoff ให้เจ้าหน้าที่เข้ามาตอบแทน AI ได้\n\n"
    "\tระบบพัฒนาด้วย FastAPI (Python) เป็น Backend, React + Vite + Tailwind CSS v4 เป็น Frontend, "
    "Supabase (PostgreSQL + pgvector) เป็นฐานข้อมูลและระบบ Authentication, "
    "โมเดล BAAI/bge-m3 สำหรับ Embedding (1024 มิติ), BAAI/bge-reranker-v2-m3 สำหรับ Reranking "
    "และ Ollama/Qwen2.5 สำหรับ LLM Generation "
    "ผลการทดสอบระบบพบว่าสามารถค้นหาข้อมูลจากเอกสารและตอบคำถามเป็นภาษาไทยได้อย่างถูกต้อง "
    "โดยรองรับการ Streaming คำตอบแบบ Real-time ผ่าน Server-Sent Events (SSE) "
    "และมีระบบป้องกัน Token หมดอายุแบบ 3 ชั้น เพื่อความเสถียรในการใช้งานต่อเนื่อง\n\n"
    "คำสำคัญ: แชทบอท AI, การสร้างคำตอบเสริมด้วยการค้นคืน, การค้นหาเชิงความหมาย, "
    "การประมวลผลภาษาธรรมชาติ, ระบบจัดการความรู้องค์กร"
)
set_paragraph(abstract_th, size=Pt(16))

doc.add_page_break()

# ============================================================
# Abstract (English)
# ============================================================
set_paragraph("Abstract", bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

p = doc.add_paragraph()
run = p.add_run("Title\t\t\t")
run.bold = True
run.font.name = FONT_EN
run.font.size = FONT_SIZE
run2 = p.add_run("SUNDAE — Enterprise AI Chatbot Platform with\n\t\t\tRetrieval-Augmented Generation")
run2.font.name = FONT_EN
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("Student\t\t\t")
run.bold = True
run.font.name = FONT_EN
run.font.size = FONT_SIZE
run2 = p.add_run("[Student Name 1]  [ID]\n\t\t\t[Student Name 2]  [ID]\n\t\t\t[Student Name 3]  [ID]")
run2.font.name = FONT_EN
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("Advisor\t\t\t")
run.bold = True
run.font.name = FONT_EN
run.font.size = FONT_SIZE
run2 = p.add_run("[Advisor Name]")
run2.font.name = FONT_EN
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("Level of Study\t\t")
run.bold = True
run.font.name = FONT_EN
run.font.size = FONT_SIZE
run2 = p.add_run("Bachelor of Science")
run2.font.name = FONT_EN
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("Major\t\t\t")
run.bold = True
run.font.name = FONT_EN
run.font.size = FONT_SIZE
run2 = p.add_run("[Department Name]")
run2.font.name = FONT_EN
run2.font.size = FONT_SIZE

p = doc.add_paragraph()
run = p.add_run("Academic Year\t\t")
run.bold = True
run.font.name = FONT_EN
run.font.size = FONT_SIZE
run2 = p.add_run("[20XX]")
run2.font.name = FONT_EN
run2.font.size = FONT_SIZE

add_empty_lines(1)

abstract_en = (
    "\tSUNDAE (Enterprise AI Chatbot Platform with Retrieval-Augmented Generation) "
    "was developed to create an intelligent chatbot system capable of providing accurate responses "
    "based on organizational knowledge bases. The system employs Retrieval-Augmented Generation (RAG), "
    "combining semantic search using vector embeddings with answer generation through a Large Language Model (LLM). "
    "Primarily designed for the Thai language, the platform enables document uploads (PDF/DOCX) "
    "for knowledge base construction, features multi-tenant user management with Role-Based Access Control "
    "(Admin, Support, User), supports omnichannel communication via Web Chat and LINE, "
    "and includes a Human Handoff system allowing human agents to intervene in AI conversations.\n\n"
    "\tThe system was built using FastAPI (Python) for the backend, React + Vite + Tailwind CSS v4 "
    "for the frontend, Supabase (PostgreSQL + pgvector) for database and authentication, "
    "BAAI/bge-m3 for text embedding (1024 dimensions), BAAI/bge-reranker-v2-m3 for result reranking, "
    "and Ollama/Qwen2.5 for LLM generation. Testing results demonstrated that the system can accurately "
    "retrieve information from documents and generate Thai-language responses in real-time "
    "via Server-Sent Events (SSE) streaming, with a three-layer token protection mechanism "
    "ensuring session stability during continuous usage.\n\n"
    "Keywords: AI Chatbot, Retrieval-Augmented Generation, Semantic Search, "
    "Natural Language Processing, Enterprise Knowledge Management"
)
set_paragraph(abstract_en, size=Pt(16))

doc.add_page_break()

# ============================================================
# กิตติกรรมประกาศ
# ============================================================
set_paragraph("กิตติกรรมประกาศ", bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

acknowledgement = (
    "\tทีมผู้พัฒนาขอกราบขอบพระคุณ [ตำแหน่ง+ชื่ออาจารย์ที่ปรึกษา] อาจารย์ที่ปรึกษาโครงงาน "
    "ที่ได้ให้คำแนะนำ แนวทาง และข้อเสนอแนะที่เป็นประโยชน์ต่อการดำเนินโครงงาน "
    "แพลตฟอร์มแชทบอท AI สำหรับองค์กร SUNDAE ตลอดระยะเวลาที่ผ่านมา "
    "ทำให้ทีมผู้พัฒนาสามารถพัฒนาและปรับปรุงผลงานให้สมบูรณ์มากยิ่งขึ้น\n\n"
    "\tขอขอบคุณ [ชื่อกรรมการ 1] และ [ชื่อกรรมการ 2] ผู้ซึ่งเป็นคณะกรรมการสอบโครงงาน "
    "ที่ได้ให้คำแนะนำและข้อคิดเห็นอันมีค่า ซึ่งช่วยให้ทีมผู้พัฒนาได้เรียนรู้และนำมาพัฒนาโครงงาน"
    "ให้มีคุณภาพ\n\n"
    "\tขอขอบคุณเพื่อนร่วมชั้นเรียนและผู้ให้ข้อมูลทุกท่าน ที่สละเวลาในการทดลองใช้งานระบบ "
    "และให้ข้อเสนอแนะที่เป็นประโยชน์ต่อการปรับปรุงระบบ\n\n"
    "\tสุดท้ายนี้ ขอขอบคุณครอบครัวที่คอยให้กำลังใจและสนับสนุนอย่างต่อเนื่อง "
    "จนทำให้ทีมผู้พัฒนามีกำลังใจในการทำโครงงานฉบับนี้จนสำเร็จลุล่วง"
)
set_paragraph(acknowledgement, size=Pt(16))

add_empty_lines(2)

set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 1]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.RIGHT)
set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 2]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.RIGHT)
set_paragraph("[ชื่อ-นามสกุล นักศึกษาคนที่ 3]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.RIGHT)
set_paragraph("[วันที่ เดือน ปี]", size=Pt(16), align=WD_ALIGN_PARAGRAPH.RIGHT)

doc.add_page_break()

# ============================================================
# สารบัญ
# ============================================================
set_paragraph("สารบัญ", bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

toc = [
    ("บทคัดย่อ", ""),
    ("Abstract", ""),
    ("กิตติกรรมประกาศ", ""),
    ("สารบัญ", ""),
    ("สารบัญตาราง", ""),
    ("สารบัญภาพ", ""),
    ("", ""),
    ("บทที่", ""),
    ("1. บทนำ", ""),
    ("   1.1 ความเป็นมาของโครงงาน", ""),
    ("   1.2 วัตถุประสงค์ของโครงงาน", ""),
    ("   1.3 ขอบเขตของโครงงาน", ""),
    ("   1.4 ซอฟต์แวร์และฮาร์ดแวร์ที่ใช้พัฒนา", ""),
    ("   1.5 ประโยชน์ที่คาดว่าจะได้รับ", ""),
    ("   1.6 ระยะเวลาที่ใช้ดำเนินโครงงาน", ""),
    ("2. ทฤษฎีและความรู้ต่าง ๆ ที่ใช้ประกอบในการทำโครงงาน", ""),
    ("   2.1 Retrieval-Augmented Generation (RAG)", ""),
    ("   2.2 Vector Embedding และ Semantic Search", ""),
    ("   2.3 Large Language Model (LLM)", ""),
    ("   2.4 FastAPI Framework", ""),
    ("   2.5 React และ Vite", ""),
    ("   2.6 Supabase และ PostgreSQL", ""),
    ("   2.7 pgvector Extension", ""),
    ("   2.8 Tailwind CSS v4", ""),
    ("   2.9 Docker และ Containerization", ""),
    ("   2.10 ระบบงานที่เกี่ยวข้อง", ""),
    ("3. การออกแบบและวิเคราะห์ระบบ", ""),
    ("   3.1 สถาปัตยกรรมระบบ (System Architecture)", ""),
    ("   3.2 แผนภาพยูสเคส (Use Case Diagram)", ""),
    ("   3.3 Use Case Description", ""),
    ("   3.4 ER Diagram", ""),
    ("   3.5 การออกแบบฐานข้อมูล", ""),
    ("   3.6 การออกแบบ API Endpoints", ""),
    ("   3.7 การออกแบบหน้าจอ (UI Design)", ""),
    ("4. ผลการดำเนินงาน", ""),
    ("   4.1 การพัฒนาระบบ", ""),
    ("   4.2 คู่มือการใช้งานระบบ", ""),
    ("   4.3 ผลการทดสอบระบบ", ""),
    ("5. สรุปผลการดำเนินงาน", ""),
    ("   5.1 สรุปผลการดำเนินงาน", ""),
    ("   5.2 ข้อจำกัดในการดำเนินงานของระบบ", ""),
    ("   5.3 แนวทางการพัฒนาต่อในอนาคต", ""),
    ("บรรณานุกรม", ""),
    ("ภาคผนวก", ""),
    ("ประวัติผู้ร่วมโครงงาน", ""),
]

for item, page in toc:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = FONT_TH
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# สารบัญตาราง
# ============================================================
set_paragraph("สารบัญตาราง", bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

table_list = [
    "1.1 ซอฟต์แวร์ที่ใช้พัฒนา",
    "1.2 ฮาร์ดแวร์ที่ใช้พัฒนา",
    "1.3 ตารางระยะเวลาดำเนินงาน",
    "3.1 Use Case: สมัครสมาชิก",
    "3.2 Use Case: เข้าสู่ระบบ",
    "3.3 Use Case: อัปโหลดเอกสาร",
    "3.4 Use Case: สนทนากับแชทบอท",
    "3.5 Use Case: จัดการ Bot",
    "3.6 Use Case: Human Handoff",
    "3.7 Use Case: อนุมัติผู้ใช้",
    "3.8 โครงสร้างตาราง organizations",
    "3.9 โครงสร้างตาราง user_profiles",
    "3.10 โครงสร้างตาราง bots",
    "3.11 โครงสร้างตาราง documents",
    "3.12 โครงสร้างตาราง document_parent_chunks",
    "3.13 โครงสร้างตาราง document_child_chunks",
    "3.14 โครงสร้างตาราง chat_sessions",
    "3.15 โครงสร้างตาราง chat_messages",
    "3.16 API Endpoints — Document Router",
    "3.17 API Endpoints — Chat Router",
    "3.18 API Endpoints — Inbox Router",
    "3.19 API Endpoints — Bot Router",
    "4.1 ผลการทดสอบฟังก์ชันการทำงาน",
]
for t in table_list:
    p = doc.add_paragraph()
    run = p.add_run(f"ตารางที่ {t}")
    run.font.name = FONT_TH
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# สารบัญภาพ
# ============================================================
set_paragraph("สารบัญภาพ", bold=True, size=Pt(18), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

fig_list = [
    "2.1 สถาปัตยกรรมของระบบ RAG",
    "2.2 กระบวนการ Embedding และ Vector Search",
    "2.3 โครงสร้างการทำงานของ FastAPI",
    "2.4 React Component Lifecycle",
    "2.5 Supabase Architecture",
    "3.1 System Architecture ของ SUNDAE",
    "3.2 แผนภาพยูสเคส (Use Case Diagram)",
    "3.3 ER Diagram ของระบบ",
    "3.4 Data Flow: การสนทนากับแชทบอท",
    "3.5 Data Flow: การอัปโหลดเอกสาร",
    "3.6 Wireframe: หน้า Login",
    "3.7 Wireframe: หน้า Dashboard",
    "3.8 Wireframe: หน้า Knowledge Base",
    "3.9 Wireframe: หน้า Web Chat",
    "3.10 Wireframe: หน้า Inbox",
    "3.11 Wireframe: หน้า Approvals",
    "4.1 หน้าเข้าสู่ระบบ",
    "4.2 หน้าสมัครสมาชิก",
    "4.3 หน้า Dashboard (Admin)",
    "4.4 หน้า Dashboard (User — รอการอนุมัติ)",
    "4.5 หน้า Knowledge Base",
    "4.6 หน้าอัปโหลดเอกสาร",
    "4.7 หน้า Web Chat",
    "4.8 หน้า Web Chat — การ Streaming คำตอบ",
    "4.9 หน้า Bots — รายการ Bot",
    "4.10 หน้า Inbox — รายการ Session",
    "4.11 หน้า Inbox — ดูข้อความและตอบกลับ",
    "4.12 หน้า Approvals — อนุมัติผู้ใช้",
]
for f in fig_list:
    p = doc.add_paragraph()
    run = p.add_run(f"ภาพที่ {f}")
    run.font.name = FONT_TH
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# บทที่ 1 — บทนำ
# ============================================================
set_paragraph("บทที่ 1", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
set_paragraph("บทนำ", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

set_heading("1.1 ความเป็นมาของโครงงาน", level=2)

intro = (
    "\tในปัจจุบัน องค์กรต่าง ๆ ทั้งภาครัฐและเอกชน มีข้อมูลและเอกสารจำนวนมากที่ต้องจัดการ "
    "ไม่ว่าจะเป็นคู่มือการทำงาน ระเบียบข้อบังคับ เอกสารนโยบาย หรือคำถามที่พบบ่อย (FAQ) "
    "การค้นหาข้อมูลจากเอกสารเหล่านี้ด้วยวิธีดั้งเดิม เช่น การค้นหาด้วยคำสำคัญ (Keyword Search) "
    "มักไม่สามารถตอบสนองความต้องการของผู้ใช้ได้อย่างมีประสิทธิภาพ เนื่องจากไม่เข้าใจ "
    "บริบทและความหมายของคำถาม (Gao et al., 2024)\n\n"
    "\tเทคโนโลยี Large Language Model (LLM) เช่น GPT, Qwen, LLaMA "
    "ได้รับความนิยมอย่างแพร่หลายในการสร้างแชทบอทอัจฉริยะ อย่างไรก็ตาม LLM มีข้อจำกัดสำคัญ "
    "คือปัญหา Hallucination — การสร้างคำตอบที่ดูสมเหตุสมผลแต่ไม่ถูกต้องตามข้อเท็จจริง "
    "(Huang et al., 2023) เทคนิค Retrieval-Augmented Generation (RAG) จึงถูกพัฒนาขึ้นเพื่อ "
    "แก้ไขปัญหานี้ โดยการค้นหาข้อมูลที่เกี่ยวข้องจากฐานความรู้ก่อน แล้วจึงส่งให้ LLM "
    "สร้างคำตอบจากข้อมูลที่ค้นพบ ทำให้คำตอบมีความถูกต้องและอ้างอิงได้ (Lewis et al., 2020)\n\n"
    "\tสำหรับภาษาไทย การใช้งาน RAG มีความท้าทายเพิ่มเติม เนื่องจากภาษาไทยไม่มีช่องว่าง "
    "ระหว่างคำ ทำให้การตัดคำ (Word Segmentation) และการสร้าง Embedding ต้องใช้โมเดลที่"
    "รองรับภาษาไทยโดยเฉพาะ (Lowphansirikul et al., 2021)\n\n"
    "\tดังนั้น โครงงานนี้จึงได้พัฒนาแพลตฟอร์ม SUNDAE ซึ่งเป็นระบบแชทบอท AI แบบ RAG "
    "สำหรับองค์กร ที่รองรับภาษาไทยเป็นหลัก สามารถอัปโหลดเอกสาร สร้างฐานความรู้อัตโนมัติ "
    "และตอบคำถามผู้ใช้ได้อย่างแม่นยำ พร้อมระบบจัดการผู้ใช้ การอนุมัติ และ Human Handoff "
    "เพื่อรองรับการใช้งานในระดับองค์กรจริง"
)
set_paragraph(intro, size=Pt(16))

set_heading("1.2 วัตถุประสงค์ของโครงงาน", level=2)
objectives = [
    "1. เพื่อพัฒนาแพลตฟอร์มแชทบอท AI ที่ใช้เทคนิค RAG ในการตอบคำถามจากฐานความรู้ขององค์กร",
    "2. เพื่อพัฒนาระบบจัดการเอกสารและฐานความรู้ที่รองรับภาษาไทย",
    "3. เพื่อพัฒนาระบบจัดการผู้ใช้แบบ Multi-tenant พร้อม Role-Based Access Control",
    "4. เพื่อศึกษาประสิทธิภาพของระบบ RAG ในการตอบคำถามภาษาไทย",
]
for obj in objectives:
    set_paragraph(obj, size=Pt(16), indent_left=1)

set_heading("1.3 ขอบเขตของโครงงาน", level=2)

set_paragraph("โครงงานนี้มีขอบเขตการพัฒนาระบบ โดยแบ่งประเภทผู้ใช้งานออกเป็น 3 ประเภท ดังนี้", size=Pt(16))

set_paragraph("1.3.1 ผู้ใช้งานทั่วไป (User)", bold=True, size=Pt(16), indent_left=1)
user_scope = [
    "1. สมัครสมาชิกและเข้าสู่ระบบด้วยอีเมลและรหัสผ่าน",
    "2. สนทนากับแชทบอท AI ผ่าน Web Chat",
    "3. รับคำตอบแบบ Streaming (Real-time)",
    "4. ดูประวัติการสนทนา",
    "5. แก้ไขข้อมูลบัญชีผู้ใช้",
    "6. รอการอนุมัติจาก Admin/Support ก่อนใช้งานระบบ",
]
for s in user_scope:
    set_paragraph(s, size=Pt(16), indent_left=2)

set_paragraph("1.3.2 เจ้าหน้าที่ Support", bold=True, size=Pt(16), indent_left=1)
support_scope = [
    "1. อนุมัติหรือปฏิเสธการสมัครสมาชิกของผู้ใช้",
    "2. จัดการ Bot — สร้าง แก้ไข ลบ",
    "3. อัปโหลดเอกสารเข้าฐานความรู้ (PDF/DOCX)",
    "4. ดูรายการ Chat Session ใน Inbox",
    "5. เข้าไปตอบกลับผู้ใช้แทน AI (Human Handoff)",
    "6. เปลี่ยนสถานะ Session (active → human_takeover → resolved)",
]
for s in support_scope:
    set_paragraph(s, size=Pt(16), indent_left=2)

set_paragraph("1.3.3 ผู้ดูแลระบบ (Admin)", bold=True, size=Pt(16), indent_left=1)
admin_scope = [
    "1. มีสิทธิ์ทั้งหมดของ Support",
    "2. ลบ Bot",
    "3. จัดการข้อมูลผู้ใช้ทั้งหมดในองค์กร",
    "4. ดูสถิติการใช้งานระบบใน Dashboard",
    "5. กำหนดสิทธิ์ (Role) ของผู้ใช้",
]
for s in admin_scope:
    set_paragraph(s, size=Pt(16), indent_left=2)

set_heading("1.4 ซอฟต์แวร์และฮาร์ดแวร์ที่ใช้พัฒนา", level=2)

set_paragraph("1. ซอฟต์แวร์", bold=True, size=Pt(16), indent_left=1)
add_table(
    ["ลำดับ", "ซอฟต์แวร์", "หน้าที่"],
    [
        ["1", "Visual Studio Code", "เครื่องมือพัฒนาโค้ด (IDE)"],
        ["2", "Python 3.11+", "ภาษาหลักสำหรับ Backend"],
        ["3", "Node.js 18+", "Runtime สำหรับ Frontend"],
        ["4", "FastAPI", "Web Framework สำหรับ REST API"],
        ["5", "React 18 + Vite", "Frontend Framework + Build Tool"],
        ["6", "Tailwind CSS v4", "CSS Utility Framework"],
        ["7", "Supabase", "Database + Authentication Platform"],
        ["8", "Ollama", "Local LLM Inference Server"],
        ["9", "Docker", "Containerization Platform"],
        ["10", "Git + GitHub", "Version Control"],
        ["11", "Figma", "UI/UX Design Tool"],
        ["12", "Draw.io", "Diagram Design Tool"],
    ],
)

set_paragraph("2. ฮาร์ดแวร์", bold=True, size=Pt(16), indent_left=1)
add_table(
    ["รายการ", "สเปค"],
    [
        ["CPU", "[ระบุ CPU]"],
        ["RAM", "16 GB ขึ้นไป (สำหรับ Embedding + LLM)"],
        ["Storage", "SSD 256 GB ขึ้นไป"],
        ["GPU", "[ระบุ GPU — ถ้ามี]"],
        ["OS", "Windows 11"],
    ],
)

set_paragraph("ภาษาที่ใช้ในการพัฒนา", bold=True, size=Pt(16), indent_left=1)
langs = [
    "1. Python — Backend API, AI Services, Text Processing",
    "2. TypeScript — Frontend (React Components, State Management)",
    "3. SQL — Database Schema, Migrations, RPC Functions",
    "4. HTML/CSS — Frontend Layout + Styling (Tailwind CSS)",
]
for l in langs:
    set_paragraph(l, size=Pt(16), indent_left=2)

set_heading("1.5 ประโยชน์ที่คาดว่าจะได้รับ", level=2)
benefits = [
    "1. ได้ระบบแชทบอท AI ที่สามารถตอบคำถามจากฐานความรู้ขององค์กรได้อย่างแม่นยำ",
    "2. ลดภาระงานของเจ้าหน้าที่ในการตอบคำถามซ้ำ ๆ",
    "3. ผู้ใช้สามารถเข้าถึงข้อมูลได้ตลอด 24 ชั่วโมง ผ่าน Web Chat",
    "4. ได้เรียนรู้เทคโนโลยี RAG, Vector Search, LLM ที่เป็นเทรนด์ปัจจุบัน",
    "5. ได้ประสบการณ์พัฒนาระบบ Full-stack ตั้งแต่ Frontend, Backend, Database จนถึง AI/ML",
    "6. สามารถนำไปต่อยอดเป็นผลิตภัณฑ์สำหรับองค์กรจริงได้",
]
for b in benefits:
    set_paragraph(b, size=Pt(16), indent_left=1)

set_heading("1.6 ระยะเวลาที่ใช้ดำเนินโครงงาน", level=2)
add_table(
    ["ลำดับ", "กิจกรรม", "เดือนที่ 1", "เดือนที่ 2", "เดือนที่ 3", "เดือนที่ 4"],
    [
        ["1", "ศึกษาและวิเคราะห์ความต้องการ", "/", "", "", ""],
        ["2", "ออกแบบระบบ (Architecture + DB + UI)", "/", "/", "", ""],
        ["3", "พัฒนา Backend (FastAPI + AI Services)", "", "/", "/", ""],
        ["4", "พัฒนา Frontend (React + Tailwind)", "", "/", "/", ""],
        ["5", "พัฒนาระบบ Auth + RLS Security", "", "", "/", ""],
        ["6", "ทดสอบและแก้ไขข้อผิดพลาด", "", "", "/", "/"],
        ["7", "จัดทำเอกสาร", "", "", "", "/"],
        ["8", "นำเสนอโครงงาน", "", "", "", "/"],
    ],
)

doc.add_page_break()

# ============================================================
# บทที่ 2 — ทฤษฎี
# ============================================================
set_paragraph("บทที่ 2", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
set_paragraph("ทฤษฎีและความรู้ต่าง ๆ ที่ใช้ประกอบในการทำโครงงาน", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

set_paragraph(
    "\tในบทนี้จะกล่าวถึงทฤษฎี เทคโนโลยี และเครื่องมือต่าง ๆ ที่ใช้ในการพัฒนาโครงงาน "
    "แพลตฟอร์มแชทบอท AI สำหรับองค์กร SUNDAE โดยครอบคลุมตั้งแต่แนวคิดพื้นฐานของ "
    "Retrieval-Augmented Generation ไปจนถึงเทคโนโลยีต่าง ๆ ที่นำมาใช้งาน",
    size=Pt(16)
)

set_heading("2.1 Retrieval-Augmented Generation (RAG)", level=2)
rag_text = (
    "\tRetrieval-Augmented Generation (RAG) เป็นเทคนิคที่ผสมผสานการค้นคืนข้อมูล (Information Retrieval) "
    "เข้ากับการสร้างข้อความ (Text Generation) โดย LLM (Lewis et al., 2020) "
    "แนวคิดหลักคือการค้นหาข้อมูลที่เกี่ยวข้องจากฐานความรู้ก่อน แล้วจึงส่งข้อมูลเหล่านั้น "
    "เป็น Context ให้ LLM ใช้ในการสร้างคำตอบ ทำให้:\n\n"
    "\t1. ลดปัญหา Hallucination — คำตอบอ้างอิงจากข้อมูลจริง\n"
    "\t2. อัปเดตความรู้ได้ — เพียงเพิ่มเอกสารใหม่โดยไม่ต้อง Fine-tune โมเดล\n"
    "\t3. ตรวจสอบได้ — สามารถอ้างอิงแหล่งที่มาของคำตอบได้\n\n"
    "\tRAG Pipeline ในโครงงานนี้ประกอบด้วย:\n"
    "\t- Document Ingestion: อัปโหลดและแบ่ง Chunk ด้วย Thai Text Splitter\n"
    "\t- Embedding: แปลง Chunk เป็น Vector ด้วย BGE-M3\n"
    "\t- Retrieval: ค้นหา Chunk ที่เกี่ยวข้องด้วย Cosine Similarity (pgvector)\n"
    "\t- Reranking: จัดอันดับผลลัพธ์ใหม่ด้วย BGE-Reranker-v2-M3\n"
    "\t- Generation: สร้างคำตอบด้วย Ollama/Qwen2.5"
)
set_paragraph(rag_text, size=Pt(16))

set_heading("2.2 Vector Embedding และ Semantic Search", level=2)
embedding_text = (
    "\tVector Embedding คือกระบวนการแปลงข้อความ (Text) เป็นเวกเตอร์ตัวเลขในปริภูมิหลายมิติ "
    "(High-Dimensional Space) โดยข้อความที่มีความหมายใกล้เคียงกันจะมีค่า Vector ที่อยู่ใกล้กัน "
    "(Reimers & Gurevych, 2019)\n\n"
    "\tโครงงานนี้ใช้โมเดล BAAI/bge-m3 ซึ่งเป็น Multilingual Embedding Model "
    "ที่รองรับภาษาไทยและมี 1024 มิติ ร่วมกับ pgvector Extension ของ PostgreSQL "
    "ในการทำ Cosine Similarity Search\n\n"
    "\tReranker (BAAI/bge-reranker-v2-m3) เป็น Cross-Encoder ที่ใช้จัดอันดับผลลัพธ์ "
    "จาก Vector Search ใหม่ โดยพิจารณาทั้ง Query และ Document พร้อมกัน "
    "ทำให้ได้ผลลัพธ์ที่แม่นยำกว่า Bi-Encoder เพียงอย่างเดียว"
)
set_paragraph(embedding_text, size=Pt(16))

set_heading("2.3 Large Language Model (LLM)", level=2)
llm_text = (
    "\tLarge Language Model (LLM) คือโมเดล AI ขนาดใหญ่ที่ถูกฝึกสอนด้วยข้อมูลข้อความจำนวนมหาศาล "
    "สามารถเข้าใจและสร้างภาษาธรรมชาติได้ (Brown et al., 2020)\n\n"
    "\tโครงงานนี้ใช้ Ollama เป็น Local LLM Server สำหรับ Inference โดยรองรับโมเดลหลายขนาด:\n"
    "\t- qwen2.5:3b (4 GB RAM) — ใช้ในปัจจุบัน\n"
    "\t- qwen2.5:7b (8 GB RAM) — สำหรับเครื่องที่มี RAM เพียงพอ\n"
    "\t- qwen3:14b (16 GB RAM) — คุณภาพสูงสุด\n\n"
    "\tข้อดีของการใช้ Local LLM:\n"
    "\t1. Data Privacy — ข้อมูลไม่ถูกส่งออกนอกองค์กร\n"
    "\t2. ไม่มีค่าใช้จ่าย API\n"
    "\t3. ปรับแต่ง keep_alive, temperature ได้ตามต้องการ"
)
set_paragraph(llm_text, size=Pt(16))

set_heading("2.4 FastAPI Framework", level=2)
fastapi_text = (
    "\tFastAPI เป็น Web Framework ของ Python ที่ออกแบบมาสำหรับสร้าง API "
    "มีความเร็วสูง รองรับ Async/Await และมี Auto-generated API Documentation (Swagger) "
    "(Ramírez, 2018)\n\n"
    "\tในโครงงานนี้ FastAPI ถูกใช้เป็น Backend API Server โดยมี Routers สำหรับ:\n"
    "\t- Document Management (อัปโหลด, ลิสต์, ลบเอกสาร)\n"
    "\t- Chat (Omnichannel chat + SSE Streaming)\n"
    "\t- Inbox (Session Management + Human Handoff)\n"
    "\t- Bot (CRUD operations)\n"
    "\t- Health Check"
)
set_paragraph(fastapi_text, size=Pt(16))

set_heading("2.5 React และ Vite", level=2)
react_text = (
    "\tReact เป็น JavaScript Library สำหรับสร้าง User Interface "
    "โดยใช้แนวคิด Component-Based Architecture (Meta, 2013) "
    "Vite เป็น Build Tool รุ่นใหม่ที่มีความเร็วสูงกว่า Webpack "
    "ใช้ ES Modules ทำให้ Hot Module Replacement (HMR) ทำงานเร็วมาก\n\n"
    "\tในโครงงานนี้ใช้ React 18 + TypeScript + Vite + Zustand (State Management) "
    "โดยมี Pages ทั้งหมด: LoginPage, DashboardPage, WebChatPage, "
    "KnowledgeBasePage, BotsPage, InboxPage, ApprovalsPage, IntegrationPage"
)
set_paragraph(react_text, size=Pt(16))

set_heading("2.6 Supabase และ PostgreSQL", level=2)
supa_text = (
    "\tSupabase เป็น Open Source Backend-as-a-Service ที่สร้างบน PostgreSQL "
    "ให้บริการ Database, Authentication, Storage และ Realtime Subscriptions (Supabase, 2020)\n\n"
    "\tในโครงงานนี้ใช้ Supabase สำหรับ:\n"
    "\t1. PostgreSQL Database — เก็บข้อมูลทั้งหมด\n"
    "\t2. Authentication — JWT-based auth (signIn, signUp, refreshToken)\n"
    "\t3. Row Level Security (RLS) — ควบคุมการเข้าถึงข้อมูลระดับ row\n"
    "\t4. RPC Functions — match_child_chunks สำหรับ vector search"
)
set_paragraph(supa_text, size=Pt(16))

set_heading("2.7 pgvector Extension", level=2)
pgvector_text = (
    "\tpgvector เป็น Extension ของ PostgreSQL สำหรับเก็บและค้นหา Vector Embedding "
    "รองรับ Cosine Similarity, L2 Distance และ Inner Product (pgvector, 2023)\n\n"
    "\tในโครงงานนี้ใช้ pgvector เก็บ embedding vector(1024) ในตาราง document_child_chunks "
    "และใช้ RPC Function match_child_chunks ทำ Cosine Similarity Search"
)
set_paragraph(pgvector_text, size=Pt(16))

set_heading("2.8 Tailwind CSS v4", level=2)
tailwind_text = (
    "\tTailwind CSS เป็น Utility-First CSS Framework ที่ช่วยให้สามารถ Style "
    "Component ได้โดยตรงใน HTML/JSX โดยไม่ต้องเขียน CSS แยก (Wathan, 2017)\n\n"
    "\tเวอร์ชัน 4 มีการปรับปรุง Engine ใหม่ ทำให้ Build เร็วขึ้นและรองรับ "
    "CSS Nesting, Container Queries และ Dynamic Values ได้ดีขึ้น"
)
set_paragraph(tailwind_text, size=Pt(16))

set_heading("2.9 Docker และ Containerization", level=2)
docker_text = (
    "\tDocker เป็นแพลตฟอร์ม Containerization ที่ช่วยบรรจุ Application "
    "พร้อม Dependencies ทั้งหมดเข้า Container ทำให้สามารถ Deploy ได้ทุก Environment "
    "อย่างสม่ำเสมอ (Merkel, 2014)\n\n"
    "\tในโครงงานนี้ใช้ Docker Compose จัดการ 3 Services:\n"
    "\t1. Backend (FastAPI) — Port 8001\n"
    "\t2. Frontend (React + Nginx) — Port 80\n"
    "\t3. Ollama (LLM Server) — Port 11434"
)
set_paragraph(docker_text, size=Pt(16))

set_heading("2.10 ระบบงานที่เกี่ยวข้อง", level=2)
related_text = (
    "\t1. ChatGPT (OpenAI) — แชทบอท AI ที่ใช้ LLM ตอบคำถามทั่วไป "
    "แต่ไม่สามารถอ้างอิงเอกสารเฉพาะขององค์กรได้ (ต้องใช้ GPTs + Knowledge Base)\n\n"
    "\t2. Langchain — Framework สำหรับสร้าง RAG Application ด้วย Python "
    "มีความซับซ้อนและ Abstraction หลายชั้น\n\n"
    "\t3. Flowise / Dify — Low-code RAG Platform ที่ใช้ UI สร้าง Workflow "
    "แต่มีข้อจำกัดในการ Customize\n\n"
    "\tSUNDAE แตกต่างจากระบบข้างต้นโดย:\n"
    "\t- พัฒนาเอง (Custom-built) ไม่พึ่ง Langchain หรือ Framework อื่น\n"
    "\t- รองรับภาษาไทยโดยเฉพาะ (Thai Text Splitter)\n"
    "\t- มี Multi-tenant + Role-Based Access Control\n"
    "\t- มี Human Handoff ให้เจ้าหน้าที่เข้ามาตอบแทน AI\n"
    "\t- ใช้ Local LLM (Ollama) ไม่ส่งข้อมูลออกนอกองค์กร"
)
set_paragraph(related_text, size=Pt(16))

doc.add_page_break()

# ============================================================
# บทที่ 3 — การออกแบบและวิเคราะห์ระบบ
# ============================================================
set_paragraph("บทที่ 3", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
set_paragraph("การออกแบบและวิเคราะห์ระบบ", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

set_heading("3.1 สถาปัตยกรรมระบบ (System Architecture)", level=2)
arch_desc = (
    "\tระบบ SUNDAE ใช้สถาปัตยกรรมแบบ Three-Tier Architecture โดยแบ่งเป็น:\n\n"
    "\t1. Presentation Layer (Frontend) — React + Vite + Tailwind CSS v4\n"
    "\t   ทำหน้าที่แสดงผล UI และรับ Input จากผู้ใช้\n\n"
    "\t2. Application Layer (Backend) — FastAPI (Python)\n"
    "\t   ทำหน้าที่ประมวลผล Business Logic, AI Pipeline, Authentication\n\n"
    "\t3. Data Layer (Database) — Supabase PostgreSQL + pgvector\n"
    "\t   ทำหน้าที่เก็บข้อมูลทั้งหมด พร้อม RLS Security\n\n"
    "[แทรกภาพที่ 3.1 System Architecture ของ SUNDAE]"
)
set_paragraph(arch_desc, size=Pt(16))

set_heading("3.2 แผนภาพยูสเคส (Use Case Diagram)", level=2)
set_paragraph(
    "\tระบบ SUNDAE มี Actors หลัก 3 ประเภท:\n"
    "\t- User (ผู้ใช้ทั่วไป)\n"
    "\t- Support (เจ้าหน้าที่)\n"
    "\t- Admin (ผู้ดูแลระบบ)\n\n"
    "[แทรกภาพที่ 3.2 Use Case Diagram]",
    size=Pt(16)
)

set_heading("3.3 Use Case Description", level=2)

# Use Case: สมัครสมาชิก
set_paragraph("ตารางที่ 3.1 Use Case: สมัครสมาชิก", bold=True, size=Pt(14))
add_table(
    ["หัวข้อ", "รายละเอียด"],
    [
        ["Use Case Name", "สมัครสมาชิก (Register)"],
        ["Actor", "User"],
        ["Description", "ผู้ใช้สร้างบัญชีใหม่ด้วย Email, Password, Full Name"],
        ["Precondition", "ยังไม่มีบัญชีในระบบ"],
        ["Main Flow", "1. ผู้ใช้กด Tab 'สมัครสมาชิก'\n2. กรอก Email, Password, Full Name\n3. กด 'สมัครสมาชิก'\n4. ระบบสร้าง auth user + user_profiles (role=user, is_approved=false)\n5. แสดงข้อความ 'สมัครสำเร็จ กรุณาเข้าสู่ระบบ'"],
        ["Postcondition", "บัญชีถูกสร้าง รอ Admin/Support อนุมัติ"],
        ["Exception", "Email ซ้ำ → แจ้ง error"],
    ],
)

# Use Case: สนทนากับแชทบอท
set_paragraph("ตารางที่ 3.4 Use Case: สนทนากับแชทบอท", bold=True, size=Pt(14))
add_table(
    ["หัวข้อ", "รายละเอียด"],
    [
        ["Use Case Name", "สนทนากับแชทบอท (Chat with AI)"],
        ["Actor", "User (approved)"],
        ["Description", "ผู้ใช้ส่งข้อความ ระบบ RAG ค้นหาข้อมูลและสร้างคำตอบ"],
        ["Precondition", "ผู้ใช้ login แล้ว + is_approved = true"],
        ["Main Flow", "1. ผู้ใช้เลือก Bot\n2. พิมพ์ข้อความ\n3. ระบบ Embed ข้อความ → Vector Search → Rerank → LLM Generate\n4. คำตอบ Stream แบบ Real-time ผ่าน SSE\n5. บันทึกทั้ง query + answer ลง chat_messages"],
        ["Postcondition", "ข้อความและคำตอบถูกบันทึก"],
        ["Exception", "ไม่พบ Bot → แจ้ง error\nToken หมดอายุ → auto-refresh"],
    ],
)

# Use Case: Human Handoff
set_paragraph("ตารางที่ 3.6 Use Case: Human Handoff", bold=True, size=Pt(14))
add_table(
    ["หัวข้อ", "รายละเอียด"],
    [
        ["Use Case Name", "Human Handoff (เจ้าหน้าที่เข้าตอบแทน AI)"],
        ["Actor", "Support / Admin"],
        ["Description", "เจ้าหน้าที่ดูข้อความใน Inbox และตอบกลับผู้ใช้แทน AI"],
        ["Precondition", "มี Chat Session ที่ต้องการ Human Handoff"],
        ["Main Flow", "1. เจ้าหน้าที่เปิดหน้า Inbox\n2. เลือก Session\n3. เปลี่ยนสถานะเป็น 'human_takeover'\n4. พิมพ์ข้อความตอบกลับ (role=admin)\n5. เปลี่ยนสถานะเป็น 'resolved' เมื่อจบ"],
        ["Postcondition", "ข้อความของเจ้าหน้าที่ถูกบันทึก + session status อัปเดต"],
    ],
)

set_heading("3.4 ER Diagram", level=2)
set_paragraph(
    "[แทรกภาพที่ 3.3 ER Diagram ของระบบ]\n\n"
    "ระบบ SUNDAE มีตารางหลัก 7 ตาราง ที่เชื่อมโยงกันผ่าน Foreign Key:",
    size=Pt(16)
)

set_heading("3.5 การออกแบบฐานข้อมูล", level=2)

set_paragraph("ตารางที่ 3.8 โครงสร้างตาราง organizations", bold=True, size=Pt(14))
add_table(
    ["Column", "Type", "Constraint", "Description"],
    [
        ["id", "UUID", "PK, DEFAULT uuid_generate_v4()", "รหัสองค์กร"],
        ["name", "TEXT", "NOT NULL", "ชื่อองค์กร"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "วันที่สร้าง"],
    ],
)

set_paragraph("ตารางที่ 3.9 โครงสร้างตาราง user_profiles", bold=True, size=Pt(14))
add_table(
    ["Column", "Type", "Constraint", "Description"],
    [
        ["id", "UUID", "PK, FK → auth.users", "รหัสผู้ใช้"],
        ["email", "TEXT", "", "อีเมล"],
        ["full_name", "TEXT", "", "ชื่อ-นามสกุล"],
        ["role", "TEXT", "DEFAULT 'user'", "บทบาท (user/support/admin)"],
        ["is_approved", "BOOLEAN", "DEFAULT false", "สถานะการอนุมัติ"],
        ["organization_id", "UUID", "FK → organizations", "รหัสองค์กร"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "วันที่สร้าง"],
    ],
)

set_paragraph("ตารางที่ 3.10 โครงสร้างตาราง bots", bold=True, size=Pt(14))
add_table(
    ["Column", "Type", "Constraint", "Description"],
    [
        ["id", "UUID", "PK", "รหัส Bot"],
        ["name", "TEXT", "NOT NULL", "ชื่อ Bot"],
        ["prompt", "TEXT", "", "System Prompt"],
        ["organization_id", "UUID", "FK → organizations", "รหัสองค์กร"],
        ["line_access_token", "TEXT", "", "LINE Access Token"],
        ["is_web_enabled", "BOOLEAN", "DEFAULT true", "เปิด Web Chat"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "วันที่สร้าง"],
    ],
)

set_paragraph("ตารางที่ 3.13 โครงสร้างตาราง document_child_chunks", bold=True, size=Pt(14))
add_table(
    ["Column", "Type", "Constraint", "Description"],
    [
        ["id", "UUID", "PK", "รหัส Chunk"],
        ["parent_chunk_id", "UUID", "FK → parent_chunks", "รหัส Parent Chunk"],
        ["document_id", "UUID", "FK → documents", "รหัสเอกสาร"],
        ["content", "TEXT", "", "เนื้อหา Chunk"],
        ["embedding", "VECTOR(1024)", "", "Vector Embedding"],
        ["chunk_index", "INTEGER", "", "ลำดับ Chunk"],
        ["created_at", "TIMESTAMPTZ", "DEFAULT now()", "วันที่สร้าง"],
    ],
)

set_heading("3.6 การออกแบบ API Endpoints", level=2)

set_paragraph("ตารางที่ 3.16 API Endpoints — Document Router", bold=True, size=Pt(14))
add_table(
    ["Method", "Endpoint", "Auth Required", "Description"],
    [
        ["POST", "/documents/upload", "Admin/Support", "อัปโหลดเอกสาร (PDF/DOCX)"],
        ["GET", "/documents/", "Approved User", "รายการเอกสาร"],
        ["DELETE", "/documents/{id}", "Admin/Support", "ลบเอกสารและ Chunks"],
    ],
)

set_paragraph("ตารางที่ 3.17 API Endpoints — Chat Router", bold=True, size=Pt(14))
add_table(
    ["Method", "Endpoint", "Auth Required", "Description"],
    [
        ["POST", "/chat/ask", "Approved User", "ส่งข้อความ รับคำตอบ"],
        ["POST", "/chat/ask-stream", "Approved User", "ส่งข้อความ รับ Streaming SSE"],
    ],
)

set_paragraph("ตารางที่ 3.18 API Endpoints — Inbox Router", bold=True, size=Pt(14))
add_table(
    ["Method", "Endpoint", "Auth Required", "Description"],
    [
        ["GET", "/inbox/sessions", "Admin/Support", "รายการ Chat Sessions"],
        ["GET", "/inbox/sessions/{id}/messages", "Admin/Support", "ข้อความใน Session"],
        ["POST", "/inbox/sessions/{id}/reply", "Admin/Support", "ตอบกลับ (Human Handoff)"],
        ["PATCH", "/inbox/sessions/{id}/status", "Admin/Support", "เปลี่ยนสถานะ Session"],
    ],
)

set_paragraph("ตารางที่ 3.19 API Endpoints — Bot Router", bold=True, size=Pt(14))
add_table(
    ["Method", "Endpoint", "Auth Required", "Description"],
    [
        ["GET", "/bots/", "Approved User", "รายการ Bots"],
        ["POST", "/bots/", "Admin/Support", "สร้าง Bot ใหม่"],
        ["PUT", "/bots/{id}", "Admin/Support", "แก้ไข Bot"],
        ["DELETE", "/bots/{id}", "Admin", "ลบ Bot"],
    ],
)

set_heading("3.7 การออกแบบหน้าจอ (UI Design)", level=2)
set_paragraph(
    "\t[แทรกภาพที่ 3.6 - 3.11 Wireframes ของแต่ละหน้า]\n\n"
    "\tระบบ SUNDAE ออกแบบ UI ตามหลัก Material Design โดยใช้ Tailwind CSS v4 "
    "และ NT CI Design System เป็นแนวทาง มี Sidebar Navigation สำหรับเมนูหลัก "
    "และ Responsive Design รองรับการใช้งานบนหน้าจอทุกขนาด",
    size=Pt(16)
)

doc.add_page_break()

# ============================================================
# บทที่ 4 — ผลการดำเนินงาน
# ============================================================
set_paragraph("บทที่ 4", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
set_paragraph("ผลการดำเนินงาน", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

set_heading("4.1 การพัฒนาระบบ", level=2)
dev_result = (
    "\tระบบ SUNDAE ได้ถูกพัฒนาขึ้นตามการออกแบบในบทที่ 3 โดยใช้เทคโนโลยีต่าง ๆ ดังนี้:\n\n"
    "\t- Backend: FastAPI (Python) — มี 5 Routers: document, chat, inbox, bot, health\n"
    "\t- Frontend: React 18 + Vite + Tailwind CSS v4 + TypeScript — มี 8 Pages\n"
    "\t- Database: Supabase PostgreSQL + pgvector — มี 7 Tables + 9 SQL Migrations\n"
    "\t- AI Services: BGE-M3 (Embedding) + BGE-Reranker-v2-M3 + Ollama/Qwen2.5\n"
    "\t- Auth: Supabase Auth + JWT + 3-Layer Token Protection\n"
    "\t- Security: RLS Policies บนทุก Table + get_my_role() SECURITY DEFINER\n"
    "\t- Deployment: Docker Compose (3 Services)\n\n"
    "\tในระหว่างการพัฒนาพบ Bug จำนวน 11 รายการ (B1-B11) ซึ่งได้แก้ไขเรียบร้อยทั้งหมด "
    "รายละเอียดอยู่ในภาคผนวก"
)
set_paragraph(dev_result, size=Pt(16))

set_heading("4.2 คู่มือการใช้งานระบบ", level=2)

set_paragraph("4.2.1 หน้าเข้าสู่ระบบ", bold=True, size=Pt(16), indent_left=1)
set_paragraph(
    "\t[แทรกภาพที่ 4.1 หน้าเข้าสู่ระบบ]\n\n"
    "\tผู้ใช้กรอก Email และ Password แล้วกดปุ่ม 'เข้าสู่ระบบ' "
    "ระบบจะตรวจสอบข้อมูลกับ Supabase Auth หากถูกต้องจะนำไปหน้า Dashboard "
    "หากบัญชียังไม่ได้รับการอนุมัติ จะแสดงหน้า Lockout Screen",
    size=Pt(16)
)

set_paragraph("4.2.2 หน้าสมัครสมาชิก", bold=True, size=Pt(16), indent_left=1)
set_paragraph(
    "\t[แทรกภาพที่ 4.2 หน้าสมัครสมาชิก]\n\n"
    "\tผู้ใช้กดแท็บ 'สมัครสมาชิก' กรอก Email, Password, ชื่อ-นามสกุล "
    "แล้วกดปุ่ม 'สมัครสมาชิก' ระบบจะสร้างบัญชีและแจ้งให้ Login ใหม่ "
    "บัญชีจะมีสถานะ is_approved = false จนกว่า Admin/Support จะอนุมัติ",
    size=Pt(16)
)

set_paragraph("4.2.3 หน้า Dashboard", bold=True, size=Pt(16), indent_left=1)
set_paragraph(
    "\t[แทรกภาพที่ 4.3, 4.4 หน้า Dashboard]\n\n"
    "\tDashboard แสดงผลตาม Role:\n"
    "\t- Admin: แสดงสถิติจำนวน Bot, เอกสาร, Session, ผู้ใช้\n"
    "\t- Support: แสดง Session ที่รอดำเนินการ\n"
    "\t- User (approved): แสดงข้อมูลทั่วไป\n"
    "\t- User (not approved): แสดง Lockout Screen",
    size=Pt(16)
)

set_paragraph("4.2.4 หน้า Knowledge Base", bold=True, size=Pt(16), indent_left=1)
set_paragraph(
    "\t[แทรกภาพที่ 4.5, 4.6]\n\n"
    "\tAdmin/Support สามารถอัปโหลดเอกสาร PDF/DOCX ได้ "
    "ระบบจะทำ Chunking → Embedding อัตโนมัติ "
    "แสดงรายการเอกสาร + สถานะ (processing/completed/error) "
    "สามารถลบเอกสารและ Chunks ทั้งหมดได้",
    size=Pt(16)
)

set_paragraph("4.2.5 หน้า Web Chat", bold=True, size=Pt(16), indent_left=1)
set_paragraph(
    "\t[แทรกภาพที่ 4.7, 4.8]\n\n"
    "\tผู้ใช้เลือก Bot → พิมพ์ข้อความ → กด Send "
    "คำตอบจาก AI จะ Stream แบบ Real-time ทีละตัวอักษร ผ่าน SSE "
    "มีปุ่ม Cancel สำหรับหยุด Stream ระหว่างรอคำตอบ",
    size=Pt(16)
)

set_paragraph("4.2.6 หน้า Inbox (Human Handoff)", bold=True, size=Pt(16), indent_left=1)
set_paragraph(
    "\t[แทรกภาพที่ 4.10, 4.11]\n\n"
    "\tAdmin/Support ดูรายการ Session ทั้งหมด เลือก Session เพื่อดูข้อความ "
    "Layout เป็น Admin Perspective:\n"
    "\t- ข้อความลูกค้า (role=user) → ชิดซ้าย สีขาว\n"
    "\t- ข้อความ AI (role=assistant) → ชิดขวา สีเหลือง\n"
    "\t- ข้อความเจ้าหน้าที่ (role=admin) → ชิดขวา สีน้ำเงิน\n"
    "\tสามารถเปลี่ยนสถานะ Session และพิมพ์ตอบกลับได้",
    size=Pt(16)
)

set_paragraph("4.2.7 หน้า Approvals", bold=True, size=Pt(16), indent_left=1)
set_paragraph(
    "\t[แทรกภาพที่ 4.12]\n\n"
    "\tAdmin/Support ดูรายชื่อผู้ใช้ที่รอการอนุมัติ "
    "กดปุ่ม 'อนุมัติ' เพื่อเปลี่ยน is_approved = true "
    "ระบบตรวจสอบว่า update สำเร็จจริงหรือไม่ด้วย .select() หลัง .update()",
    size=Pt(16)
)

set_heading("4.3 ผลการทดสอบระบบ", level=2)

set_paragraph(
    "\tทดสอบฟังก์ชันการทำงานหลักของระบบ โดยผลการทดสอบแสดงดังตาราง",
    size=Pt(16)
)

set_paragraph("ตารางที่ 4.1 ผลการทดสอบฟังก์ชันการทำงาน", bold=True, size=Pt(14))
add_table(
    ["ลำดับ", "ฟังก์ชัน", "ผลการทดสอบ"],
    [
        ["1", "สมัครสมาชิก + auto-create profile", "ผ่าน"],
        ["2", "เข้าสู่ระบบ + Token refresh", "ผ่าน"],
        ["3", "Role-based access control (3 roles)", "ผ่าน"],
        ["4", "อัปโหลดเอกสาร + Chunking + Embedding", "ผ่าน"],
        ["5", "Chat — RAG Pipeline + Streaming SSE", "ผ่าน"],
        ["6", "Chat — Token auto-refresh ระหว่าง Stream", "ผ่าน"],
        ["7", "Inbox — ดู Sessions + Messages", "ผ่าน"],
        ["8", "Inbox — Human Handoff ตอบกลับ", "ผ่าน"],
        ["9", "Bot CRUD", "ผ่าน"],
        ["10", "Approvals — อนุมัติผู้ใช้", "ผ่าน"],
        ["11", "RLS Security — ป้องกัน unauthorized access", "ผ่าน"],
        ["12", "Ollama keep_alive 4h — ไม่ unload model", "ผ่าน"],
    ],
)

doc.add_page_break()

# ============================================================
# บทที่ 5 — สรุปผล
# ============================================================
set_paragraph("บทที่ 5", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
set_paragraph("สรุปผลการดำเนินงาน", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

set_heading("5.1 สรุปผลการดำเนินงาน", level=2)
summary = (
    "\tโครงงานแพลตฟอร์มแชทบอท AI สำหรับองค์กร SUNDAE ได้พัฒนาระบบแชทบอทอัจฉริยะ "
    "ที่ใช้เทคนิค Retrieval-Augmented Generation (RAG) ในการตอบคำถามจากฐานความรู้ขององค์กร "
    "สามารถสรุปผลได้ดังนี้:\n\n"
    "\t1. พัฒนาระบบ RAG ที่ใช้ BGE-M3 Embedding + pgvector Cosine Similarity Search "
    "+ BGE-Reranker + Ollama/Qwen2.5 LLM ในการค้นหาและสร้างคำตอบจากเอกสารที่อัปโหลด "
    "รองรับภาษาไทยเป็นหลัก\n\n"
    "\t2. พัฒนาระบบจัดการเอกสารที่รองรับ PDF/DOCX โดยใช้ Thai Text Splitter "
    "ในการแบ่ง Chunk แบบ Parent-Child Strategy เพื่อประสิทธิภาพในการค้นหา\n\n"
    "\t3. พัฒนาระบบจัดการผู้ใช้แบบ Multi-tenant พร้อม Role-Based Access Control "
    "(Admin, Support, User) และ Row Level Security (RLS) บนทุกตาราง\n\n"
    "\t4. พัฒนาระบบ Streaming SSE สำหรับแสดงคำตอบ AI แบบ Real-time "
    "พร้อม 3-Layer Token Protection เพื่อความเสถียร\n\n"
    "\t5. พัฒนาระบบ Human Handoff ให้เจ้าหน้าที่เข้ามาตอบแทน AI ผ่านหน้า Inbox\n\n"
    "\tระบบผ่านการทดสอบฟังก์ชันหลักทั้ง 12 รายการ และแก้ไข Bug จำนวน 11 รายการ "
    "จนทำงานได้อย่างเสถียร"
)
set_paragraph(summary, size=Pt(16))

set_heading("5.2 ข้อจำกัดในการดำเนินงานของระบบ", level=2)
limitations = (
    "\t1. LLM Model — ปัจจุบันใช้ qwen2.5:3b (4 GB RAM) ซึ่งมีความแม่นยำระดับปานกลาง-สูง "
    "เนื่องจากเครื่อง Dev มี RAM ไม่เพียงพอสำหรับ qwen3:14b ที่ต้องการ 16 GB\n\n"
    "\t2. LINE Integration — ยังอยู่ในสถานะ Mock เฉพาะ UI "
    "ยังไม่ได้สร้าง Webhook Endpoint และ LINE Reply Service จริง\n\n"
    "\t3. Evaluation — ยังไม่ได้ทำ Quantitative Evaluation เช่น BLEU, ROUGE, F1 Score "
    "สำหรับวัดคุณภาพคำตอบของ RAG Pipeline\n\n"
    "\t4. Scalability — ยังไม่ได้ทดสอบ Load Test สำหรับ Concurrent Users จำนวนมาก\n\n"
    "\t5. File Types — รองรับเฉพาะ PDF และ DOCX ยังไม่รองรับ Excel, PowerPoint, รูปภาพ"
)
set_paragraph(limitations, size=Pt(16))

set_heading("5.3 แนวทางการพัฒนาต่อในอนาคต", level=2)
future = (
    "\t1. LINE Webhook Integration — สร้าง POST /webhook/line endpoint "
    "เชื่อมต่อ RAG Pipeline กับ LINE Messaging API ให้ลูกค้าสนทนาผ่าน LINE ได้\n\n"
    "\t2. Upgrade LLM — อัปเกรดเป็น qwen2.5:7b หรือ qwen3:14b "
    "เมื่อมี Hardware ที่เหมาะสม เพื่อคุณภาพคำตอบที่ดีขึ้น\n\n"
    "\t3. Multi-modal RAG — รองรับเอกสารที่มีรูปภาพ ตาราง กราฟ "
    "โดยใช้ Vision Model ช่วยสกัดข้อมูล\n\n"
    "\t4. Analytics Dashboard — แสดงสถิติการใช้งาน เช่น จำนวนคำถาม "
    "คะแนนความพึงพอใจ หัวข้อที่ถามบ่อย\n\n"
    "\t5. Fine-tuning — Fine-tune Embedding Model และ LLM "
    "ด้วยข้อมูลเฉพาะขององค์กร เพื่อเพิ่มความแม่นยำ\n\n"
    "\t6. Dark Mode — เพิ่ม Theme Switcher สำหรับ UI\n\n"
    "\t7. User Profile Edit — ให้ผู้ใช้แก้ไข Full Name ของตัวเอง\n\n"
    "\t8. Load Testing — ทดสอบ Concurrent Users เพื่อรองรับการ Deploy Production"
)
set_paragraph(future, size=Pt(16))

doc.add_page_break()

# ============================================================
# บรรณานุกรม
# ============================================================
set_paragraph("บรรณานุกรม", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

references = [
    "Brown, T. B., et al. (2020). Language models are few-shot learners. Advances in Neural Information Processing Systems, 33, 1877-1901.",
    "Gao, Y., et al. (2024). Retrieval-augmented generation for large language models: A survey. arXiv preprint arXiv:2312.10997.",
    "Huang, L., et al. (2023). A survey on hallucination in large language models. arXiv preprint arXiv:2311.05232.",
    "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459-9474.",
    "Lowphansirikul, T., et al. (2021). WangchanBERTa: Pretraining transformer-based Thai language models. arXiv preprint arXiv:2101.09635.",
    "Merkel, D. (2014). Docker: Lightweight Linux containers for consistent development and deployment. Linux Journal, 2014(239), 2.",
    "Meta. (2013). React: A JavaScript library for building user interfaces. https://react.dev/",
    "pgvector. (2023). Open-source vector similarity search for Postgres. https://github.com/pgvector/pgvector",
    "Ramírez, S. (2018). FastAPI framework, high performance, easy to learn, fast to code. https://fastapi.tiangolo.com/",
    "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. Proceedings of EMNLP-IJCNLP, 3982-3992.",
    "Supabase. (2020). The open source Firebase alternative. https://supabase.com/",
    "Wathan, A. (2017). Tailwind CSS: A utility-first CSS framework. https://tailwindcss.com/",
]

for ref in references:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = FONT_TH
    run.font.size = Pt(16)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ============================================================
# ภาคผนวก
# ============================================================
set_paragraph("ภาคผนวก", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

set_heading("ภาคผนวก ก — SQL Migrations", level=2)
add_table(
    ["#", "Migration", "สถานะ", "วัตถุประสงค์"],
    [
        ["001", "Main Schema", "รันแล้ว", "สร้าง Tables + match_child_chunks RPC"],
        ["002", "Add Columns", "รันแล้ว", "เพิ่ม LINE + Session fields"],
        ["003", "user_profiles + RLS", "รันแล้ว", "Role system + get_my_role()"],
        ["004", "Auth Trigger", "รันแล้ว", "Auto-create profile on signup"],
        ["005", "Support Account", "รันแล้ว", "สร้าง Support user"],
        ["006", "Bot Filter", "รันแล้ว", "Vector search filter by bot_id"],
        ["007", "Admin Role", "รันแล้ว", "Allow 'admin' in chat_messages"],
        ["008", "Fix Org RLS", "รันแล้ว", "Fix 406 on organizations"],
        ["009", "Fix UPDATE RLS", "รันแล้ว", "Add WITH CHECK to UPDATE"],
    ],
)

set_heading("ภาคผนวก ข — Bugs ที่พบและแก้ไข", level=2)
add_table(
    ["#", "Bug", "สาเหตุ", "วิธีแก้"],
    [
        ["B1", "Loading screen ค้าง", "ไม่มี .catch() + timeout", "เพิ่ม .catch() + 5s timeout"],
        ["B2", "ขาด email/full_name columns", "Schema ไม่ครบ", "เพิ่ม columns (migration 002)"],
        ["B3", "RLS block profile read", "ไม่มี SELECT policy สำหรับตัวเอง", "เพิ่ม id = auth.uid()"],
        ["B4", "RLS infinite recursion", "Subquery อ่าน user_profiles ซ้ำ", "สร้าง get_my_role() SECURITY DEFINER"],
        ["B5", "Stream ไม่ยิง request", "Dynamic import ค้างใน Vite", "เปลี่ยนเป็น static import"],
        ["B6", "Organizations 406", "RLS ใช้ JWT claim ที่ไม่มี", "เปลี่ยนเป็น subquery"],
        ["B7", "Login 'Failed to fetch'", "ไม่มี catch block", "เพิ่ม catch + Thai error"],
        ["B8", "Approve ไม่ทำงาน", "ขาด WITH CHECK", "เพิ่ม WITH CHECK + .select()"],
        ["B9", "Stream ล้มหลังใช้นาน", "Token หมดอายุไม่ refresh", "ใช้ getValidToken()"],
        ["B10", "Ollama unload model", "keep_alive 30m สั้นเกินไป", "เปลี่ยนเป็น 4h"],
        ["B11", "Inbox layout ผิด", "ใช้ User perspective", "เปลี่ยนเป็น Admin perspective"],
    ],
)

doc.add_page_break()

# ============================================================
# ประวัติผู้ร่วมโครงงาน
# ============================================================
set_paragraph("ประวัติผู้ร่วมโครงงาน", bold=True, size=Pt(20), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

for i in range(1, 4):
    set_paragraph(f"ผู้ร่วมโครงงานคนที่ {i}", bold=True, size=Pt(16))
    info = [
        f"ชื่อ-นามสกุล\t\t[ชื่อ-นามสกุล นักศึกษาคนที่ {i}]",
        f"รหัสนักศึกษา\t\t[XXXXXXXXXX]",
        f"คณะ\t\t\t[ชื่อคณะ]",
        f"ภาควิชา\t\t\t[ชื่อภาควิชา]",
        f"อีเมล\t\t\t[email@example.com]",
        f"โทรศัพท์\t\t\t[0XX-XXX-XXXX]",
    ]
    for line in info:
        set_paragraph(line, size=Pt(16), indent_left=1)
    add_empty_lines(1)

# ============================================================
# SAVE
# ============================================================
import sys
sys.stdout.reconfigure(encoding='utf-8')
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SUNDAE_Report_5_Chapters.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
